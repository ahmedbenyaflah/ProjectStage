#!/usr/bin/env python3
"""
Build a simplified, submission-ready PFE DOCX (3 chapters, incremental methodology).
Run from repo root: python3 scripts/build_pfe_3chapters_simplified_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title_block(doc: Document, title: str, subtitle_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    for line in subtitle_lines:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.add_run(line)
    doc.add_paragraph()


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)


def add_figure(doc: Document, image_path: Path, caption: str, max_width_cm: float = 16.0) -> None:
    if not image_path.is_file():
        add_body_paragraph(doc, f"[Figure manquante : {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(max_width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    out = root / "rapport_pfe_automatisation_logs_smtp_simplifie.docx"
    shots = root / "kpis screenshots"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_block(
        doc,
        "Analyse automatisée des journaux SMTP et reconstruction des parcours de messagerie",
        [
            "Rapport de projet de fin d’études",
            "Entreprise d’accueil : Orange Tunisie",
            "Spécialité : Ingénierie logicielle et données",
            "Technologies : Elastic Stack (Elasticsearch, Logstash, Kibana), FastAPI, React",
            "",
            "Réalisé par : [Nom et prénom]",
            "Encadrant entreprise : [Nom]",
            "Encadrant académique (ISAMM) : [Nom]",
            "Année universitaire : [20.. / 20..]",
        ],
    )

    doc.add_heading("Table des matières", level=1)
    toc_lines = [
        "Introduction générale .................................................... 3",
        "Chapitre 1 — Contexte général ............................................ 4",
        "    1.1 Présentation d’Orange Tunisie",
        "    1.2 Problématique et cadre du projet",
        "    1.3 Méthodologie de travail : modèle incrémental",
        "Chapitre 2 — Analyse et conception ....................................... 7",
        "    2.1 Rappels sur le protocole SMTP",
        "    2.2 Analyse des besoins (acteurs, besoins fonctionnels et non fonctionnels)",
        "    2.3 Architecture proposée",
        "Chapitre 3 — Réalisation ................................................... 11",
        "    3.1 Environnement de développement",
        "    3.2 Intégration de l’offre Elastic Stack",
        "    3.3 Développement de l’API avec FastAPI",
        "    3.4 Conception de l’interface avec React",
        "Conclusion générale ........................................................ 15",
        "Bibliographie .............................................................. 16",
        "Annexes — Captures d’écran (indicateurs et tableaux de bord) ............... 17",
    ]
    for line in toc_lines:
        doc.add_paragraph(line)
    doc.add_page_break()

    # --- Introduction générale ---
    doc.add_heading("Introduction générale", level=1)
    intro = """
Les services de messagerie d’un opérateur comme Orange Tunisie reposent sur de nombreux serveurs qui enregistrent chaque étape d’un courriel : réception, filtrage, routage, livraison ou rejet. Ces traces se présentent sous forme de fichiers journaux très volumineux, répartis sur plusieurs familles de machines et plusieurs formats de lignes. Pour l’équipe technique, retrouver « ce qui s’est passé pour un message précis » demande souvent beaucoup de recherches manuelles et du temps.

Ce mémoire présente un travail de fin d’études dont l’objectif est de proposer une chaîne logicielle capable d’analyser automatiquement ces journaux SMTP, de reconstruire le parcours d’un message (ou « mail journey ») à partir d’informations fragmentées, et de les rendre exploitables via une base de recherche (Elasticsearch), une API (FastAPI) et une application web (React), complétée par des tableaux de bord Kibana pour le suivi des indicateurs.

La problématique centrale est la suivante : comment passer d’événements dispersés dans plusieurs fichiers à une vision ordonnée et consultable du trajet d’un courriel, afin de réduire le délai de diagnostic et d’améliorer la qualité du support aux équipes internes ?

La démarche suit un modèle incrémental : à chaque cycle, une partie du système (extraction, indexation, API ou interface) est livrée, testée avec les parties déjà existantes, puis enrichie selon les retours. Cette approche limite les risques sur un périmètre aussi large (données massives, sécurité, ergonomie).

Le document est structuré en trois chapitres, conformément aux usages fréquemment attendus pour un rapport de stage à l’ISAMM en entreprise partenaire comme Orange Tunisie. Le premier chapitre situe l’entreprise, le contexte métier et la méthode. Le second chapitre rappelle les bases du protocole SMTP, formalise les besoins et décrit l’architecture retenue. Le troisième chapitre détaille la mise en œuvre technique : environnement, intégration Elastic Stack, API et interface utilisateur. Des captures d’écran issues de l’outil (indicateurs, DNSBL, erreurs, volumes) sont regroupées en annexe.
""".strip()
    for para in intro.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_page_break()

    # --- Chapitre 1 ---
    doc.add_heading("Chapitre 1 — Contexte général", level=1)

    doc.add_heading("1.1 Présentation d’Orange Tunisie", level=2)
    c11 = """
Orange Tunisie est l’une des principales filiales du groupe Orange en Afrique et au Moyen-Orient. L’entreprise propose des services de téléphonie mobile, d’accès Internet et de solutions professionnelles. Dans ce cadre, la messagerie électronique constitue un service critique : elle doit rester disponible, sécurisée et traçable.

Les équipes techniques supervisent une infrastructure complexe où les journaux serveur jouent un rôle clé pour l’audit, la conformité et le dépannage. La quantité d’informations produite chaque jour rend impossible une lecture « ligne par ligne » sans outils adaptés.
""".strip()
    for para in c11.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_heading("1.2 Problématique et cadre du projet", level=2)
    c12 = """
Les journaux SMTP d’Orange Tunisie décrivent des événements liés aux messages (acceptation, refus temporaire, refus définitif, filtrage antispam ou antivirus, etc.). En pratique, un même message peut laisser plusieurs empreintes dans différents fichiers ou sur différents serveurs, et les identifiants visibles ne sont pas toujours identiques d’une étape à l’autre. Cette fragmentation complique la reconstitution du parcours réel du courriel.

Avant la mise en place d’une plateforme dédiée, une investigation pouvait impliquer plusieurs recherches manuelles, l’ouverture de très gros fichiers texte et un risque d’erreur humaine. Les besoins exprimés par l’équipe technique peuvent se résumer ainsi : réduire le temps de diagnostic, offrir une recherche fiable par critères (adresses, plage horaire, codes d’erreur), et permettre une vision agrégée des tendances (erreurs fréquentes, volumes, réputation IP via listes DNSBL).

Le projet de fin d’études s’inscrit dans cette attente : concevoir et implémenter une solution qui transforme les journaux bruts en données structurées, indexées et interrogées depuis une application moderne, tout en s’appuyant sur l’écosystème Elastic pour l’exploration visuelle.
""".strip()
    for para in c12.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_heading("1.3 Méthodologie de travail : modèle incrémental", level=2)
    c13 = """
Le modèle incrémental consiste à découper le travail en versions successives, chacune ajoutant des fonctionnalités utilisables. Contrairement à une livraison unique en fin de projet, l’incrément privilégie des cycles courts : analyser un petit lot de besoins, concevoir la partie correspondante, l’implémenter, puis valider avec l’encadrant avant d’élargir le périmètre.

Dans ce stage, les incréments suivants ont structuré la progression :

• Incrément 1 — compréhension des journaux : lecture d’échantillons réels, identification des champs utiles (expéditeur, destinataire, codes SMTP, identifiants de session) et définition du format cible (documents JSON pour l’index Elasticsearch).

• Incrément 2 — chaîne d’ingestion : scripts Python de parsing (messages envoyés et reçus), normalisation des enregistrements et chargement dans Elasticsearch avec un schéma de « mail journey » cohérent.

• Incrément 3 — exposition des données : conception d’une API FastAPI sécurisée par authentification (JWT), points de recherche et agrégations utiles au front.

• Incrément 4 — interface opérateur : application React pour les formulaires de recherche, affichage des détails et intégration d’iframes Kibana selon le profil utilisateur.

• Incrément 5 — fiabilisation : tests manuels sur plages horaires contrôlées, ajustements de performance, fonctions complémentaires (par exemple surveillance DNSBL, alertes) selon les priorités de l’équipe.

À chaque fin d’incrément, une démonstration permettait de vérifier que le module livré fonctionnait avec le reste du système. Cette approche réduit les surprises en fin de stage et aligne le développement sur les besoins réels du terrain.
""".strip()
    for para in c13.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_page_break()

    # --- Chapitre 2 ---
    doc.add_heading("Chapitre 2 — Analyse et conception", level=1)

    doc.add_heading("2.1 Rappels sur le protocole SMTP", level=2)
    c21 = """
SMTP (Simple Mail Transfer Protocol) est le protocole standard utilisé pour transférer le courrier entre serveurs sur Internet. Un client ou un serveur émetteur envoie des commandes textuelles (MAIL FROM, RCPT TO, DATA…) et reçoit des codes numériques à trois chiffres qui indiquent le succès, un report temporaire ou un échec définitif.

Les fichiers journaux observés dans ce projet reprennent typiquement ces codes (par exemple 250 pour une acceptation, 4xx pour une erreur temporaire, 5xx pour un rejet). Comprendre SMTP permet d’interpréter correctement les messages d’erreur affichés aux opérateurs et de relier plusieurs lignes de log à une même tentative d’envoi.

Le parcours d’un message peut ainsi être vu comme une suite d’événements SMTP et de traitements associés (filtrage, file d’attente, relais). La conception logicielle doit prévoir comment regrouper ces événements en un objet logique « parcours » consultable depuis l’interface.
""".strip()
    for para in c21.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_heading("2.2 Analyse des besoins", level=2)
    doc.add_heading("2.2.1 Acteurs", level=3)
    c221 = """
On distingue principalement : l’administrateur ou ingénieur support (recherche ciblée, lecture du parcours), le responsable supervision (vision agrégée via Kibana), et le système d’information (comptes utilisateurs, droits d’accès). L’API sert d’intermédiaire contrôlé entre l’interface web et Elasticsearch.
""".strip()
    add_body_paragraph(doc, c221)

    doc.add_heading("2.2.2 Besoins fonctionnels", level=3)
    c222 = """
• Ingérer et indexer les journaux SMTP issus des parseurs métier.\n• Rechercher des messages par critères temporels et par adresses.\n• Afficher le détail d’un parcours reconstruit.\n• Proposer des tableaux de bord pour les volumes, erreurs, spam, virus et contrôles DNSBL.\n• Gérer l’authentification et, selon le profil, l’accès aux fonctions d’administration.
""".strip()
    add_body_paragraph(doc, c222)

    doc.add_heading("2.2.3 Besoins non fonctionnels", level=3)
    c223 = """
• Performance : réponses de recherche rapides sur de grands index.\n• Sécurité : authentification, limitation des plages temporelles de requête, principe du moindre privilège.\n• Disponibilité raisonnable dans un environnement de laboratoire ou de préproduction.\n• Maintenabilité : code modulaire (parseurs, construction de requêtes Elasticsearch, composants React).\n• Ergonomie : formulaires clairs et libellés compréhensibles pour l’équipe support.
""".strip()
    add_body_paragraph(doc, c223)

    doc.add_heading("2.3 Architecture proposée", level=2)
    c23 = """
L’architecture retenue suit une séparation classique en trois niveaux.

Couche données : Elasticsearch stocke les documents représentant les parcours et les événements associés. Les index peuvent être typés selon les jeux de données (envoyés, reçus, parcours DNSBL, etc.).

Couche services : une application FastAPI expose des routes REST sécurisées. Elle construit les requêtes vers Elasticsearch, applique les règles métier (validation des paramètres, filtrage) et renvoie des réponses JSON consommées par le front.

Couche présentation : une application React offre les écrans de connexion, de recherche et de consultation du détail. Des vues Kibana intégrées complètent l’outil pour les indicateurs préconfigurés.

Les parseurs Python, exécutés dans le cadre du pipeline de données, assurent la transformation des fichiers bruts vers le format d’indexation. Cette organisation facilite l’évolution indépendante de chaque bloc (par exemple changer un écran sans modifier le moteur d’indexation).
""".strip()
    for para in c23.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_page_break()

    # --- Chapitre 3 ---
    doc.add_heading("Chapitre 3 — Réalisation", level=1)

    doc.add_heading("3.1 Environnement de développement", level=2)
    c31 = """
Le développement s’est appuyé sur un poste de travail sous Linux, Python 3 pour les parseurs et l’API, Node.js pour la chaîne de build React, et des conteneurs ou services locaux pour Elasticsearch et Kibana lors des tests. Les variables d’environnement permettent d’adapter les URL (cluster Elasticsearch, instance Kibana) sans modifier le code source.

Le dépôt du projet regroupe le dossier backend_pipeline (API FastAPI, authentification, scan DNSBL, construction de requêtes) et logs_filter_frontend_elastic (application React). Cette cohabitation facilite la reprise du travail par une autre équipe.
""".strip()
    for para in c31.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_heading("3.2 Intégration de l’offre Elastic Stack", level=2)
    c32 = """
Elasticsearch sert de moteur de recherche et d’agrégation sur les documents issus des journaux. Les gabarits d’index (templates) assurent des types de champs cohérents pour les recherches par plage de dates et les filtres sur les adresses.

Kibana permet de construire des tableaux de bord : volumes de messages, répartition des codes d’erreur, tendances spam ou virus, suivis DNSBL. Dans l’application React, des URL d’intégration (mode embed) pointent vers ces tableaux de bord pour que l’utilisateur reste dans un cadre unique tout en profitant de la puissance de visualisation d’Elastic.

Logstash peut être utilisé dans des déploiements où l’on souhaite une ingestion continue depuis des files ou des bus de messages ; dans la version décrite ici, l’accent est mis sur les parseurs Python et l’API, Elasticsearch restant le point central de persistance et d’interrogation.
""".strip()
    for para in c32.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_heading("3.3 Développement de l’API avec FastAPI", level=2)
    c33 = """
FastAPI a été choisi pour sa rapidité de développement, sa documentation OpenAPI automatique et son bon support des modèles Pydantic pour valider les entrées. L’application expose notamment des routes d’authentification (connexion, inscription selon la politique de déploiement), des endpoints de recherche sur les parcours, et des utilitaires liés à la surveillance des listes noires DNS.

Le module query_builder regroupe la logique de construction des clauses Elasticsearch à partir des critères utilisateur, ce qui évite de dupliquer des fragments de requête dans chaque route. Un schéma de données pour les parcours (journey_schema) harmonise les champs provenant de sources hétérogènes.

La sécurité repose sur des jetons JWT vérifiés à chaque requête protégée, complétés par une distinction des rôles lorsque des fonctions d’administration sont activées.
""".strip()
    for para in c33.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_heading("3.4 Conception de l’interface avec React", level=2)
    c34 = """
L’interface utilisateur est une application monopage développée avec React. Les principaux écrans comprennent l’authentification, un formulaire de recherche avec contrôle de la plage horaire maximale, l’affichage des résultats et un panneau de détail pour le parcours sélectionné. Le contexte d’authentification centralise le jeton et l’état de session.

Des composants dédiés intègrent les tableaux de bord Kibana en iframe pour les profils qui y ont droit (par exemple distinctions N1/N2/N3 selon la configuration). Le fichier de configuration regroupe les URL de l’API, d’Elasticsearch lorsque le front interroge directement certaines métadonnées, et les liens d’embed Kibana.

L’objectif ergonomique est de permettre à un opérateur de passer d’une plainte client (« mon message n’est pas arrivé ») à une lecture ordonnée des étapes du message en quelques clics, lorsque les données sont présentes dans les index.
""".strip()
    for para in c34.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_page_break()

    doc.add_heading("Conclusion générale", level=1)
    conc = """
Ce mémoire a présenté une solution d’analyse automatisée des journaux SMTP et de reconstruction des parcours de messagerie, réalisée au sein d’Orange Tunisie dans une logique Big Data orientée recherche et visualisation.

Les apports principaux sont la structuration des données brutes en documents exploitables dans Elasticsearch, l’exposition sécurisée via une API FastAPI, et une interface React intégrant les indicateurs Kibana. Le modèle incrémental a permis de valider progressivement chaque brique avec l’équipe encadrante.

Les perspectives naturelles consistent à élargir les sources prises en charge, à affiner encore les règles de corrélation entre lignes de log pour les cas limites, et à renforcer l’automatisation des alertes métiers. Ce travail illustre comment une chaîne moderne (Elastic Stack + API + SPA) peut répondre concrètement aux contraintes de volume et de délai rencontrées en exploitation messagerie.
""".strip()
    for para in conc.split("\n\n"):
        add_body_paragraph(doc, para)

    doc.add_page_break()

    doc.add_heading("Bibliographie", level=1)
    bib = """
[1] Elastic. Documentation Elasticsearch. https://www.elastic.co/guide/en/elasticsearch/reference/current/index.html\n[2] Elastic. Documentation Kibana. https://www.elastic.co/guide/en/kibana/current/index.html\n[3] FastAPI. Documentation officielle. https://fastapi.tiangolo.com/\n[4] J. Klensin. RFC 5321 — Simple Mail Transfer Protocol. IETF, 2008.\n[5] React. Documentation. https://react.dev/\n[6] Fielding, R. T. Architectural Styles and the Design of Network-based Software Architectures (REST). Thèse de doctorat, UC Irvine, 2000.\n[7] CHESS study group. DNS-based Blackhole List (DNSBL) — usages et bonnes pratiques opérationnelles (ressources communautaires et fournisseurs de listes).
""".strip()
    add_body_paragraph(doc, bib)

    doc.add_page_break()

    doc.add_heading("Annexes — Captures d’écran (indicateurs et tableaux de bord)", level=1)
    add_body_paragraph(
        doc,
        "Les figures suivantes illustrent l’outil tel qu’utilisé pour le suivi des volumes, des erreurs SMTP, du filtrage antispam/antivirus et des contrôles DNSBL. Les légendes reprennent les noms des fichiers sources.",
    )

    # Deterministic order: overview first, then thematic groups
    image_specs = [
        ("dashboard overview.png", "Figure A.1 — Vue d’ensemble du tableau de bord."),
        ("log men bara .png", "Figure A.2 — Consultation des journaux / parcours (interface)."),
        ("sent vs recieved mails .png", "Figure A.3 — Messages envoyés versus reçus."),
        ("top senders and recipents.png", "Figure A.4 — Principaux expéditeurs et destinataires."),
        ("top domain names.png", "Figure A.5 — Domaines les plus représentés."),
        ("most common error codes .png", "Figure A.6 — Codes d’erreur les plus fréquents."),
        ("example of error code 564.png", "Figure A.7 — Exemple autour du code d’erreur 564."),
        ("example of error code 572.png", "Figure A.8 — Exemple autour du code d’erreur 572."),
        ("spam mails overtime.png", "Figure A.9 — Évolution du spam dans le temps."),
        ("virus mails overtime.png", "Figure A.10 — Évolution des messages virus dans le temps."),
        ("dnsbl check .png", "Figure A.11 — Contrôle DNSBL."),
        ("dnsbl kpis.png", "Figure A.12 — Indicateurs DNSBL."),
        ("DNSBL checks overti,e.png", "Figure A.13 — Contrôles DNSBL dans le temps."),
        ("audit.png", "Figure A.14 — Vue d’audit."),
        ("some kpis.png", "Figure A.15 — Jeu d’indicateurs (KPI)."),
        ("some more kpis.png", "Figure A.16 — Indicateurs complémentaires."),
        ("unkown transmission error log men bara .png", "Figure A.17 — Erreur de transmission inconnue (journal)."),
    ]

    for fname, caption in image_specs:
        add_figure(doc, shots / fname, caption)

    doc.save(str(out))
    print("Wrote:", out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
